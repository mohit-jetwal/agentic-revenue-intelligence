"""Generate the Project Bible: this platform, written as interview material.

Deliberately matches the layout and styling of ``GenAI-Interview-Bible.docx`` -
same chapter shape (Mental Model, How It Works, tagged Interview Questions,
Common Wrong Answers), same type scale, same palette. The reason is practical
rather than aesthetic: two documents read back to back during preparation should
not make the reader re-learn where to look.

The content is generated from a script rather than typed into Word so that it is
diffable, reviewable and reproducible. When a number in the platform changes,
this file is where the document changes with it.

Run with::

    uv run python docs/generate_project_bible.py

Every figure below is measured and traceable to a step in the build. Nothing in
this document is illustrative.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path("docs/Agentic_Revenue_Intelligence_Project_Bible.docx")

# Palette lifted from the GenAI Interview Bible so the two sit together.
INK = RGBColor(0x1A, 0x1A, 0x1A)
NAVY = RGBColor(0x08, 0x3D, 0x6B)
BLUE = RGBColor(0x0B, 0x5F, 0xA5)
GREY = RGBColor(0x5F, 0x6C, 0x7B)
LEAD = RGBColor(0x33, 0x3D, 0x47)
QUOTE = RGBColor(0x4A, 0x3B, 0x10)
DANGER = RGBColor(0x9B, 0x2C, 0x2C)
GOOD = RGBColor(0x1E, 0x6B, 0x3A)

BODY_FONT = "Georgia"
HEAD_FONT = "Segoe UI"
MONO_FONT = "Consolas"


# --------------------------------------------------------------------------
# Styling
# --------------------------------------------------------------------------


def build_styles(doc: Document) -> None:
    """Recreate the Bible's type scale."""
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.22

    for name, size, colour in (
        ("Heading 1", 21.0, NAVY),
        ("Heading 2", 15.5, BLUE),
        ("Heading 3", 12.5, NAVY),
    ):
        style = doc.styles[name]
        style.font.name = HEAD_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.italic = False
        style.font.color.rgb = colour

    doc.styles["Heading 1"].paragraph_format.space_before = Pt(20)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(9)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(15)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(6)
    doc.styles["Heading 3"].paragraph_format.space_before = Pt(12)
    doc.styles["Heading 3"].paragraph_format.space_after = Pt(4)

    from docx.enum.style import WD_STYLE_TYPE

    def custom(name: str) -> object:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        return style

    lead = custom("ProjectLead")
    lead.font.size = Pt(11.5)
    lead.font.italic = True
    lead.font.color.rgb = LEAD
    lead.paragraph_format.space_after = Pt(12)

    caption = custom("ProjectCaption")
    caption.font.name = HEAD_FONT
    caption.font.size = Pt(8.5)
    caption.font.italic = True
    caption.font.color.rgb = GREY
    caption.paragraph_format.space_after = Pt(10)

    quote = custom("ProjectQuote")
    quote.font.size = Pt(10)
    quote.font.color.rgb = QUOTE
    quote.paragraph_format.left_indent = Pt(15.85)
    quote.paragraph_format.space_before = Pt(6)
    quote.paragraph_format.space_after = Pt(6)

    bullet = custom("ProjectBullet")
    bullet.paragraph_format.left_indent = Pt(18)
    bullet.paragraph_format.space_after = Pt(3)

    code = custom("ProjectCode")
    code.font.name = MONO_FONT
    code.font.size = Pt(8.5)
    code.paragraph_format.left_indent = Pt(12)
    code.paragraph_format.space_after = Pt(2)
    code.paragraph_format.line_spacing = 1.0

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)


# --------------------------------------------------------------------------
# Building blocks
# --------------------------------------------------------------------------


def heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level)


def para(doc: Document, text: str, style: str | None = None) -> None:
    """A paragraph with **bold** and `code` spans rendered inline.

    Nesting is handled one level deep, because ``**a `code` span**`` is the
    natural way to write a sentence here and a single-pass split would match the
    bold first, emit its whole interior as one run, and leave the backticks
    visible as literal characters.
    """
    import re

    paragraph = doc.add_paragraph(style=style)

    def add_code(chunk: str, *, bold: bool = False) -> None:
        run = paragraph.add_run(chunk)
        run.font.name = MONO_FONT
        run.font.size = Pt(9)
        run.font.color.rgb = NAVY
        run.bold = bold

    for chunk in re.split(r"(\*\*.+?\*\*|`.+?`)", text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            for inner in re.split(r"(`.+?`)", chunk[2:-2]):
                if not inner:
                    continue
                if inner.startswith("`") and inner.endswith("`"):
                    add_code(inner[1:-1], bold=True)
                else:
                    paragraph.add_run(inner).bold = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            add_code(chunk[1:-1])
        else:
            paragraph.add_run(chunk)


def lead(doc: Document, text: str) -> None:
    para(doc, text, "ProjectLead")


def caption(doc: Document, text: str) -> None:
    para(doc, text, "ProjectCaption")


def quote(doc: Document, text: str) -> None:
    para(doc, text, "ProjectQuote")


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        para(doc, f"•  {item}", "ProjectBullet")


def code(doc: Document, text: str) -> None:
    for line in text.strip("\n").splitlines():
        doc.add_paragraph(line or " ", style="ProjectCode")


def flag(doc: Document, text: str, colour: RGBColor = DANGER) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = colour


def table(doc: Document, rows: list[list[str]], widths: list[float] | None = None) -> None:
    """A grid table with a bold header row."""
    grid = doc.add_table(rows=len(rows), cols=len(rows[0]))
    grid.style = "Table Grid"
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = grid.cell(r, c)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(value)
            run.font.size = Pt(9)
            run.font.name = BODY_FONT
            if r == 0:
                run.bold = True
                run.font.color.rgb = NAVY
    if widths:
        for row in grid.rows:
            for cell, width in zip(row.cells, widths, strict=False):
                cell.width = Inches(width)
    doc.add_paragraph()


def question(doc: Document, tag: str, text: str) -> None:
    heading(doc, f"[{tag}]  {text}", 3)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def toc(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-2" \\h \\z \\u')
    run._r.addnext(fld)


def build(output: Path = OUTPUT) -> Path:
    """Assemble the document and write it."""
    # Imported here rather than at module scope: the content module imports the
    # helpers above, so a top-level import would be circular.
    from docs.project_bible_content import CHAPTERS, front_matter

    doc = Document()
    build_styles(doc)
    front_matter(doc)
    for chapter in CHAPTERS:
        chapter(doc)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return output


if __name__ == "__main__":
    written = build()
    print(f"wrote {written}")
    print(
        "Open in Word and run: right-click the Contents table -> "
        "Update Field -> Update entire table."
    )


__all__ = ["build", "build_styles"]
