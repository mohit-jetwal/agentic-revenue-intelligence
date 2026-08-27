"""Generate the project handbook as a Word document.

    uv run python docs/generate_handbook.py

Produces ``docs/Agentic_Revenue_Intelligence_Handbook.docx`` - a single
end-to-end reference covering the architecture, every completed step, the
defects found along the way, and how to run and extend the platform.

Written as a generator rather than hand-authored so it regenerates from one
source of truth as later steps land. Diagrams come from
``generate_handbook_diagrams.py`` and must be produced first.
"""

from __future__ import annotations

import datetime as dt
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
IMAGES = Path(__file__).parent / "images"
OUTPUT = Path(__file__).parent / "Agentic_Revenue_Intelligence_Handbook.docx"

ACCENT = RGBColor(0x1B, 0x49, 0x65)
MUTED = RGBColor(0x5A, 0x64, 0x72)
DANGER = RGBColor(0xB3, 0x40, 0x3A)
GOOD = RGBColor(0x2E, 0x7D, 0x5B)


# ---------------------------------------------------------------------------
# Building blocks
# ---------------------------------------------------------------------------


def code(document: Document, text: str) -> None:
    """A monospaced block with a light background."""
    for line in text.strip("\n").split("\n"):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.3)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        run = paragraph.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1B, 0x24, 0x30)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F4F6F9")
        paragraph._p.get_or_add_pPr().append(shading)
    document.add_paragraph().paragraph_format.space_after = Pt(4)


def callout(document: Document, label: str, text: str, colour: RGBColor = ACCENT) -> None:
    """A labelled aside - used for decisions, warnings and findings."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(8)

    tag = paragraph.add_run(f"{label}  ")
    tag.bold = True
    tag.font.size = Pt(9.5)
    tag.font.color.rgb = colour

    _rich(paragraph, text, 10)

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F7F9FC")
    paragraph._p.get_or_add_pPr().append(shading)


def para(document: Document, text: str, *, size: float = 10.5, italic: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    # Routed through the rich renderer like bullets and tables are. Adding the
    # text as one plain run leaves literal ** and backticks on the page, which
    # is how the first generated draft came out.
    _rich(paragraph, text, size)
    if italic:
        for run in paragraph.runs:
            run.italic = True


def bullets(document: Document, items: list[str], *, numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        paragraph = document.add_paragraph(style=style)
        paragraph.paragraph_format.space_after = Pt(3)
        _rich(paragraph, item, 10)


def _rich(paragraph, text: str, size: float) -> None:
    """Render ``**bold**`` and ``` `code` ``` spans, including code inside bold.

    Two passes rather than one alternation. A single pattern cannot express
    nesting, and the content genuinely nests - several bullets lead with a bold
    clause that names a function. A one-pass renderer emitted those as bold text
    containing literal backticks.
    """
    import re

    for chunk in re.split(r"(\*\*.+?\*\*)", text):
        if not chunk:
            continue
        bold = chunk.startswith("**") and chunk.endswith("**")
        inner = chunk[2:-2] if bold else chunk

        for token in re.split(r"(`[^`]+`)", inner):
            if not token:
                continue
            mono = token.startswith("`") and token.endswith("`")
            run = paragraph.add_run(token[1:-1] if mono else token)
            run.bold = bold
            run.font.size = Pt(size - 1) if mono else Pt(size)
            if mono:
                run.font.name = "Consolas"


def table(document: Document, headers: list[str], rows: list[list[str]],
          *, widths: list[float] | None = None) -> None:
    grid = document.add_table(rows=1, cols=len(headers))
    grid.style = "Light Grid Accent 1"
    grid.alignment = WD_TABLE_ALIGNMENT.CENTER

    for index, heading in enumerate(headers):
        cell = grid.rows[0].cells[index]
        cell.text = ""
        run = cell.paragraphs[0].add_run(heading)
        run.bold = True
        run.font.size = Pt(9)

    for row in rows:
        cells = grid.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = ""
            paragraph = cells[index].paragraphs[0]
            _rich(paragraph, str(value), 9)
            for run in paragraph.runs:
                run.font.size = Pt(9)

    if widths:
        for row in grid.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)

    document.add_paragraph().paragraph_format.space_after = Pt(6)


def figure(document: Document, name: str, caption: str, width: float = 6.3) -> None:
    path = IMAGES / f"{name}.png"
    if not path.is_file():
        para(document, f"[missing figure: {name}]", italic=True)
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))

    label = document.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = label.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    label.paragraph_format.space_after = Pt(12)


def heading(document: Document, text: str, level: int) -> None:
    item = document.add_heading(text, level=level)
    for run in item.runs:
        run.font.color.rgb = ACCENT
    item.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
    item.paragraph_format.space_after = Pt(6)


def page_break(document: Document) -> None:
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def toc(document: Document) -> None:
    """Insert a table-of-contents field.

    Word populates it on open (or on F9). It cannot be pre-rendered from here,
    so the placeholder text tells the reader how to refresh it rather than
    leaving an empty page that looks broken.
    """
    paragraph = document.add_paragraph()
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = r'TOC \o "1-3" \h \z \u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click here and choose 'Update Field' to build the contents."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for element in (begin, instruction, separate, placeholder, end):
        run._r.append(element)


def git_sha() -> str:
    try:
        return subprocess.run(
            ["git", "rev-parse", "--short", "HEAD"],
            cwd=ROOT, capture_output=True, text=True, check=True, timeout=10,
        ).stdout.strip()
    except Exception:  # noqa: BLE001 - provenance is best-effort
        return "unknown"


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------


def cover(document: Document) -> None:
    for _ in range(5):
        document.add_paragraph()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Agentic Revenue, Pricing\n& Promotion Intelligence")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = ACCENT

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Engineering Handbook")
    run.font.size = Pt(17)
    run.font.color.rgb = MUTED

    document.add_paragraph()
    strap = document.add_paragraph()
    strap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = strap.add_run(
        "A CPG / Retail platform where Claude and LangGraph orchestrate reasoning\n"
        "and deterministic models produce every number."
    )
    run.italic = True
    run.font.size = Pt(11)

    for _ in range(6):
        document.add_paragraph()

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Stage 1, Steps 1-5 complete\n"
        f"Generated {dt.date.today().isoformat()}  |  commit {git_sha()}\n"
        f"534 tests passing  |  ruff, mypy, bandit clean"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED

    page_break(document)


def how_to_read(document: Document) -> None:
    heading(document, "How to read this handbook", 1)
    para(document,
         "This document is the single reference for the platform: what it does, why each "
         "decision was made, what went wrong along the way, and how to run and extend it. "
         "It is organised so you can enter at any level.")

    table(document,
          ["If you want to...", "Read"],
          [
              ["Understand the shape of the system in ten minutes",
               "Part I - Orientation"],
              ["Know why the data is synthetic and what that buys",
               "Part II, Step 2"],
              ["Understand how leakage is prevented structurally",
               "Part II Step 3, and Part IV"],
              ["Understand the models and their measured results",
               "Part III - Steps 4 and 5"],
              ["Run the platform or retrain a model",
               "Part V - Operations"],
              ["Learn from the mistakes rather than repeat them",
               "Part VI - the defect catalogue"],
              ["Prepare to discuss this in an interview",
               "Part VI - talking points"],
          ],
          widths=[3.4, 2.6])

    callout(document, "A note on tone.",
            "Numbers in this document are measured, not estimated. Where something is "
            "uncertain, underspecified, or was got wrong and later corrected, it says so. "
            "The defect catalogue in Part VI is deliberately unflattering - it is the most "
            "useful section in the document.")
    page_break(document)


def part_orientation(document: Document) -> None:
    heading(document, "Part I  -  Orientation", 1)

    heading(document, "1. What this platform is", 2)
    para(document,
         "A commercial analytics platform for consumer packaged goods and retail. It answers "
         "questions a category manager actually asks - how much will we sell, did that "
         "promotion work, what should this product cost - and it answers them with a "
         "specific division of labour:")
    bullets(document, [
        "**Claude and LangGraph orchestrate reasoning.** They decide which question is "
        "being asked, which analyses to run, in what order, and how to explain the result.",
        "**Deterministic models produce every number.** No figure in any answer originates "
        "from a language model. Every quantity is traceable to a fitted model, a dataset "
        "version and a code commit.",
    ])
    callout(document, "The governing constraint.",
            "The LLM must never invent a numerical result. This is not a stylistic "
            "preference - it is what makes the platform auditable. An agent that can "
            "estimate a number is an agent whose output cannot be checked.")

    heading(document, "2. The two-stage architecture", 2)
    para(document,
         "The platform is built twice over: once locally, where iteration is fast and "
         "everything is inspectable, and once on Databricks, where it would run in "
         "production. The second build is a redeployment rather than a rewrite, and that "
         "claim is enforced by three seams rather than asserted in a document.")
    figure(document, "01_two_stage_architecture",
           "Figure 1  -  Only three components differ between local and production.")
    para(document,
         "Everything above the seams - services, models, feature logic, the tool contract, "
         "the agent layer - is identical in both columns. A model trained locally moves to "
         "Stage 2 without a code change, because it never learns where its data came from.")
    callout(document, "No AWS.",
            "This platform is Databricks-native in production. AWS is not used anywhere - "
            "no S3, no ECR, no Lambda, no managed database. The constraint is deliberate "
            "and is reflected in the migration design in Part V.")

    heading(document, "3. Layered architecture", 2)
    para(document,
         "Within either column, dependencies point strictly downward. A model may use the "
         "feature layer; the feature layer may not know a model exists.")
    figure(document, "02_layered_architecture",
           "Figure 2  -  Seven layers. The two amber layers are where implementations swap.")

    heading(document, "4. Technology stack", 2)
    table(document,
          ["Concern", "Stage 1 (local)", "Stage 2 (production)"],
          [
              ["Language", "Python 3.12+, managed by uv", "same"],
              ["Analytics store", "DuckDB over Parquet", "Databricks SQL over Delta"],
              ["Application state", "SQLite", "Databricks SQL"],
              ["Feature engineering", "pandas", "PySpark"],
              ["Models", "LightGBM, XGBoost, scikit-learn, statsmodels", "same"],
              ["Optimisation", "OR-Tools", "same"],
              ["Experiment tracking", "MLflow on SQLite", "Databricks MLflow"],
              ["Model registry", "local MLflow registry", "Unity Catalog"],
              ["Vector store", "Chroma", "Databricks Vector Search"],
              ["Reasoning", "Claude (Anthropic API)", "same"],
              ["Orchestration", "LangGraph", "same"],
              ["Contracts", "Pydantic, Pandera", "same"],
              ["API", "FastAPI", "Databricks Model Serving"],
              ["Quality gates", "ruff, mypy, bandit, pytest", "same"],
          ],
          widths=[1.7, 2.3, 2.0])

    heading(document, "5. Repository map", 2)
    code(document, """
app/            application layer
  config/       Settings - one class per section, env-prefixed
  services/     Container (DI), baseline_service, forecast_service
  tools/        AnalyticalTool ABC, ToolResult envelope, registry
  schemas/      Pydantic contracts: domain, tool_contract, baseline, forecast
  observability/  structlog logging, trace context, metrics
  api/          FastAPI app and routes

data/           data layer
  generation/   the synthetic data simulator (Step 2)
  repositories/ DataRepository ABC, Local + Databricks, PointInTimeView
  contracts/    Pandera table schemas

features/       feature layer (Step 3)
  contracts/    catalogue, specs, config loader, feature versioning
  engineering/  panel helpers, demand, pricing, promotion, inventory, temporal
  datasets/     per-model dataset builders
  repositories/ FeatureRepository ABC + local implementation

ml/             model layer
  base.py       AnalyticalModel ABC, ModelMetadata, error types
  baseline/     Step 4 - baseline sales
  forecasting/  Step 5 - demand forecasting
  <others>/     interface-only stubs for Steps 6-11

tests/          534 tests, marked by concern
configs/        YAML: dataset profiles, feature selection, model config
scripts/        generate_data, train_baseline, train_forecast
docs/           this handbook, model cards, migration designs
notebooks/      exploration and validation, one folder per step
""")
    page_break(document)


def part_foundations(document: Document) -> None:
    heading(document, "Part II  -  Foundations", 1)

    # -- Step 1 ------------------------------------------------------------
    heading(document, "Step 1  -  Skeleton and the three seams", 2)
    para(document,
         "Step 1 built no features and trained no models. It established the boundaries that "
         "make the remaining nineteen steps possible, and the reason it comes first is that "
         "these boundaries are almost impossible to retrofit.")

    heading(document, "Seam 1  -  DataRepository", 3)
    para(document,
         "An abstract base class defining every read the platform performs. Business logic "
         "receives a repository and never asks which implementation it got.")
    code(document, """
class DataRepository(ABC):
    @abstractmethod
    def get_sales(self, *, product_ids=None, store_ids=None,
                  start_date=None, end_date=None, as_of_date=None,
                  max_rows=None) -> pd.DataFrame: ...

    def as_of(self, as_of_date: date) -> PointInTimeView:
        return PointInTimeView(self, as_of_date)
""")

    heading(document, "Seam 2  -  the ToolResult envelope", 3)
    para(document,
         "Every analytical capability returns the same shape. `AnalyticalTool.run()` is "
         "marked `@final`, so a subclass cannot bypass validation, timing, tracing or error "
         "wrapping - it implements only `_execute`.")
    table(document,
          ["Field", "Purpose"],
          [
              ["`status`", "success / partial / error / invalid_input / timeout"],
              ["`model_name`, `model_version`, `dataset_version`",
               "Provenance - which artifact produced this number"],
              ["`result`", "The payload, as plain JSON"],
              ["`confidence`", "Measured, or absent. Never invented"],
              ["`assumptions`", "What the caller must know to use the number responsibly"],
              ["`warnings`", "Caveats; a non-empty list downgrades status to `partial`"],
              ["`error`", "Code, message, and a `recoverable` flag for agent re-planning"],
              ["`trace_id`", "Ties the result to its log lines"],
          ],
          widths=[2.2, 3.8])

    heading(document, "Seam 3  -  the DI container", 3)
    para(document,
         "One class with cached properties, keyed on `APP__ENVIRONMENT`. Components are "
         "built lazily, so constructing the container never requires a trained model or an "
         "API key - a missing artifact surfaces at the point of use, not at boot.")

    callout(document, "Design decision.",
            "Not a DI framework. dependency-injector and friends buy scoping, auto-wiring "
            "and declarative overrides; with four factories and one switch they would add a "
            "dependency and a mental model in exchange for nothing.")

    heading(document, "Defects found in Step 1", 3)
    table(document,
          ["Symptom", "Cause", "Fix"],
          [
              ["Settings failed to construct",
               "`**_SECTION_CONFIG` collided with an explicit `env_prefix` keyword",
               "Standalone dict constant, spread per section"],
              ["46 test errors, `AttributeError` on every log call",
               "`structlog.stdlib.add_logger_name` with `PrintLoggerFactory`",
               "Switched to `structlog.stdlib.LoggerFactory()`"],
              ["bandit B101",
               "`assert` used for type narrowing - stripped under `python -O`",
               "Explicit local-variable narrowing"],
          ],
          widths=[1.9, 2.4, 1.8])
    page_break(document)

    # -- Step 2 ------------------------------------------------------------
    heading(document, "Step 2  -  Synthetic data with hidden ground truth", 2)
    para(document,
         "The dataset is generated from a structural causal model rather than sampled from "
         "a distribution. That distinction is what makes every later step testable: because "
         "the true elasticities, promotion lifts and latent demand are known, a model can be "
         "checked against the answer rather than against its own fit.")
    figure(document, "03_data_generation",
           "Figure 3  -  Demand is simulated, then censored by inventory. Both are retained.")

    heading(document, "Scale and shape", 3)
    table(document,
          ["Property", "Value"],
          [
              ["Period", "2023-01-01 to 2025-12-31 (1,096 days)"],
              ["Products / stores", "300 / 200"],
              ["**Real product-store series**", "**6,128** - each product is listed in ~10% of stores"],
              ["Daily sales rows", "6,716,288"],
              ["Total rows across all tables", "23.6 million"],
              ["Gold tables", "11"],
          ],
          widths=[2.4, 3.6])

    heading(document, "The six deliberate confounders", 3)
    para(document,
         "A dataset where price moves at random would make elasticity trivial to recover and "
         "the platform's causal machinery pointless. These confounders are what give Steps "
         "6-8 something real to solve.")
    table(document,
          ["Confounder", "What it creates"],
          [
              ["Price endogeneity (0.45)", "Price responds to demand, biasing naive elasticity"],
              ["Commodity-cost instrument (0.55 pass-through)", "A **valid IV** - moves price, not demand"],
              ["Randomised price tests (18% of periods)", "A clean experimental subset to validate against"],
              ["Promotion targeting (0.40)", "Promotions scheduled toward seasonal peaks - selection bias"],
              ["Competitor-cost correlation (0.40)", "Competitor prices move with shared costs"],
              ["Endogenous stockouts", "Stockouts happen **because** demand spiked"],
          ],
          widths=[2.6, 3.4])

    callout(document, "The finding that matters most.",
            "Latent demand during a stockout runs 1.57x normal. Stockouts are caused by "
            "demand, not merely correlated with it. This single fact invalidated the "
            "original stockout diagnostic in Step 4 - see Part III.", DANGER)

    heading(document, "Hidden ground truth", 3)
    para(document,
         "`data/local/ground_truth/` holds the true parameters and the uncensored demand. "
         "It is **structurally unreachable** through any repository method: the local "
         "repository registers only the gold directory, so no `get_*` call and no SQL query "
         "can reach it. The single sanctioned reader is the evaluation path, which reads "
         "from disk directly.")

    heading(document, "Defects found in Step 2", 3)
    table(document,
          ["Symptom", "Cause", "Why it mattered"],
          [
              ["Stockouts were impossible - 170 days of cover",
               "Replenishment ordered against **on-hand stock, not inventory position**, "
               "stacking duplicate orders during lead time",
               "The censoring mechanism that Steps 4-5 are validated against did not exist"],
              ["Data changed between identical runs",
               "Chunk RNG used `SeedSequence.spawn()`, which mutates the parent's child counter",
               "Reproducibility - the headline claim of the step - was false"],
              ["Elasticity unidentifiable at dev scale",
               "Validation sampled the **first** partitions, giving a 3-month window",
               "A correct simulator looked broken"],
              ["Transaction table 19x too large", "Sampling rate mis-set", "Generation time and disk"],
              ["`__init__.py` files silently emptied",
               "`New-Item -ItemType File -Force` truncates existing files",
               "Packages stopped importing"],
          ],
          widths=[1.7, 2.3, 2.1])
    page_break(document)

    # -- Step 3 ------------------------------------------------------------
    heading(document, "Step 3  -  Data access, contracts and features", 2)
    para(document,
         "Step 3 is where leakage prevention became structural rather than procedural. The "
         "central insight is that the naive rule - clamp every table to the as-of date - is "
         "wrong, and being wrong in the safe direction still destroys information the "
         "platform needs.")

    heading(document, "Availability classes", 3)
    figure(document, "04_availability_classes",
           "Figure 4  -  What is knowable depends on the table, not on a global rule.")
    para(document,
         "Sales for a future date do not exist. The promotion calendar for that same date "
         "does - it was planned weeks ago. Blanket clamping would delete exactly the "
         "information a forecast depends on, and Step 5 would have been impossible.")

    heading(document, "Point-in-time correctness as a type constraint", 3)
    para(document,
         "`FeatureEngineer` requires a `PointInTimeView` and raises `TypeError` on a bare "
         "repository. Leakage through the feature layer is therefore not a discipline "
         "question - it is a construction that will not run.")

    heading(document, "The feature pipeline", 3)
    figure(document, "05_feature_pipeline",
           "Figure 5  -  The order is load-bearing, not incidental.")
    bullets(document, [
        "**Sorting precedes every shift.** A lag computed on an unsorted panel is silently wrong.",
        "**Target-derived columns are dropped last.** `revenue = units x price` recovers the "
        "target exactly, and a model given revenue learns arithmetic rather than demand.",
        "**Every temporal feature routes through `shifted_group` or `rolling_on_shifted`.** "
        "A window can never include its own row; `periods < 1` is rejected outright.",
    ])

    heading(document, "Defects found in Step 3", 3)
    table(document,
          ["Symptom", "Cause", "Fix"],
          [
              ["A query returned exactly 100,000 of 6,716,288 rows, silently",
               "Row cap applied with no signal to the caller",
               "`ResultTruncatedError` - a truncated analysis is worse than a failed one"],
              ["Near-perfect model accuracy",
               "`revenue`, `cost`, `gross_profit` reached the feature matrix",
               "`drop_target_derived()` inside the engineer"],
              ["Panels came back inexplicably sparse",
               "Products and stores sampled independently, producing never-co-listed pairs",
               "Co-listed sampling path"],
              ["`channel` parameter vanished from a method",
               "Regex-based patching of Python source",
               "Use the editor, never regex, on code"],
          ],
          widths=[1.9, 2.1, 2.0])
    page_break(document)


def part_models(document: Document) -> None:
    heading(document, "Part III  -  The models", 1)

    heading(document, "The distinction that carries both steps", 2)
    figure(document, "06_baseline_vs_forecast",
           "Figure 6  -  Same machinery, opposite counterfactual.")
    para(document,
         "A baseline estimates what **would** have happened without an intervention. A "
         "forecast predicts what **will** happen including planned interventions. Conflating "
         "them is the most common way uplift analysis goes wrong: a forecast that includes "
         "the promotion measures the promotion against itself and reports roughly zero uplift.")
    page_break(document)

    # -- Step 4 ------------------------------------------------------------
    heading(document, "Step 4  -  Baseline sales model", 2)
    para(document,
         "The question: what would sales have been under normal conditions - no promotion "
         "running, stock available? Promotion uplift is measured against it, root cause "
         "compares actual to it, and scenario simulation starts from it. A biased baseline "
         "does not fail loudly; it propagates a consistent distortion into every downstream "
         "number.")

    heading(document, "The two design decisions", 3)
    para(document, "**Promotional contamination.** Two defensible approaches, both with a real bias:")
    table(document,
          ["Approach", "Strength", "Bias"],
          [
              ["**C** - train on non-promotional rows only",
               "Clean counterfactual semantics",
               "Promotions target seasonal peaks, so dropping them under-represents high "
               "season and **overstates uplift**"],
              ["**B** - promotion features as controls, predict with them zeroed",
               "Uses all rows, avoids selection bias",
               "Asks a tree model to extrapolate to a feature combination it may rarely "
               "have seen"],
          ],
          widths=[1.9, 1.8, 2.3])
    para(document,
         "Neither argument settles it, so both were built and scored against true demand. "
         "Approach C won by 0.2 points - too narrow to generalise, and reported as such.")

    para(document, "**Stockout censoring.** Observed sales during a stockout measure "
                   "availability, not demand. Stockout rows are excluded from training, and "
                   "censored values are not lagged forward.")

    callout(document, "The finding that reshaped the step.",
            "The first full run disqualified both LightGBM candidates. The feature "
            "importances explained why: `closing_inventory_lag_1` was the single most "
            "important feature. Given inventory, the model learns that low stock predicts "
            "low sales - which is true, and is exactly the censoring relationship the step "
            "exists to avoid. Excluding stockout rows does not prevent it; the relationship "
            "is learned from partially-depleted rows and extrapolated to zero stock.", DANGER)
    table(document,
          ["LightGBM configuration", "Stockout lift", "Recovers (correct ~0.64)", "Outcome"],
          [
              ["Inventory features present", "1.12", "0.30", "**disqualified**"],
              ["Supply features excluded", "**2.48**", "**0.68**", "selected"],
          ],
          widths=[2.3, 1.3, 1.6, 1.3])
    para(document,
         "Accuracy moved only 40.4% to 40.2% WMAPE. That is the telling part: those columns "
         "were contributing censoring signal rather than demand signal, so removing them "
         "cost nothing and fixed the diagnostic outright.", italic=True)

    heading(document, "Results", 3)
    table(document,
          ["Metric", "Value", "Reading"],
          [
              ["Selected model", "`lightgbm__exclude`", "Beat the seasonal naive by 13.9 points"],
              ["Latent WMAPE", "40.4%", "Against true demand, not observed sales"],
              ["**Irreducible noise floor**", "**35.0%**", "What a model knowing the true mean would score"],
              ["**Ratio to floor**", "**1.15x**", "Most remaining error is unlearnable noise"],
              ["Stockout lift", "2.48x", "Sees through the censoring"],
              ["Interval coverage", "92.0% on 587,603 rows", "Nominal 90% - calibrated"],
              ["Backtest", "39.0% +/- 0.7%", "Stable across four expanding folds"],
          ],
          widths=[1.9, 1.8, 2.3])

    callout(document, "An unresolved tension, stated rather than resolved.",
            "The design argues bias matters more than error for a baseline, but the "
            "selection rule ranks on WMAPE - and here they disagree. LightGBM is 3.3 points "
            "more accurate; Ridge is nearly unbiased (-0.5% vs +6.7%). The over-prediction "
            "understates uplift rather than inventing it, which is the safer direction, but "
            "adding a bias term to the criterion is a real decision and belongs with Step 6.")
    page_break(document)

    # -- Step 5 ------------------------------------------------------------
    heading(document, "Step 5  -  Demand forecasting model", 2)
    para(document,
         "Step 4's baseline is a nowcast - it predicts units at date D using features at D, "
         "including yesterday's sales. That is valid for a historical counterfactual and "
         "invalid for forecasting: standing at as-of T predicting T+30, you do not know "
         "T+29's sales. Every design decision in Step 5 follows from that gap.")

    heading(document, "The horizon dataset", 3)
    figure(document, "07_horizon_dataset",
           "Figure 7  -  Each feature is placed by asking whether it is knowable at the origin.")
    table(document,
          ["Feature family", "Sourced at", "Why legitimate"],
          [
              ["lags, rollings, dynamics", "origin `t`", "`sales_daily` is OBSERVED"],
              ["competitor price and gap", "origin `t`", "OBSERVED - not knowable forward"],
              ["calendar, festival, season", "target `t+h`", "`calendar` is KNOWN_IN_ADVANCE"],
              ["planned promotion", "target `t+h`", "`promotions` is KNOWN_IN_ADVANCE"],
              ["planned price", "target `t+h`", "`pricing` is KNOWN_IN_ADVANCE"],
              ["`horizon_step`", "-", "Lets one model span every horizon"],
          ],
          widths=[2.1, 1.5, 2.4])
    para(document,
         "Target-side columns carry an `h_` prefix, so the origin/target split is visible in "
         "the feature-importance table rather than being a fact you must remember.")

    heading(document, "Why one model rather than four", 3)
    bullets(document, [
        "**Four per-horizon models on a cumulative target** - dispositive: cannot produce the "
        "daily path the interface requires, and the nested totals would be mutually incoherent.",
        "**Recursive one-step-ahead** - the killer is feature-distribution collapse, not error "
        "compounding. Feeding conditional-mean predictions forward drives rolling standard "
        "deviations and volatility toward zero, so by h=30 the model sees inputs it never saw.",
        "**Four direct models on daily targets** - 4x the compute, less data per model, no "
        "benefit once `h` is already a feature.",
    ])

    heading(document, "The embargo", 3)
    figure(document, "08_embargo_split",
           "Figure 8  -  Without the gap, a boundary training origin is scored on its own outcome.")

    heading(document, "Prediction intervals", 3)
    para(document,
         "Two things Step 4's single scalar quantile could not do. **Width must grow with "
         "horizon**, so calibration is per horizon bucket (Mondrian conformal). And **the "
         "horizon total needs its own calibration** - summing daily bounds assumes the daily "
         "errors move together perfectly, and if they were independent the sum would be too "
         "wide by roughly the square root of 90.")

    heading(document, "The training pipeline", 3)
    figure(document, "09_training_pipeline",
           "Figure 9  -  The order is the correctness argument.")

    heading(document, "Results", 3)
    table(document,
          ["Model", "WMAPE", "Bias", "Mean FVA", "Train"],
          [
              ["**xgboost**", "**43.8%**", "+8.4%", "**+12.6 pp**", "34.4s"],
              ["lightgbm", "47.5%", "+19.9%", "+9.5 pp", "11.3s"],
              ["horizon_seasonal_naive", "56.0%", "+7.2%", "-", "0.4s"],
              ["horizon_naive", "78.1%", "+36.1%", "-19.8 pp", "0.4s"],
          ],
          widths=[2.0, 1.0, 1.0, 1.1, 0.9])
    para(document,
         "800 series, 548,754 rows, 475 seconds. Forecast Value Added is **positive at every "
         "horizon bucket** (+11 to +15 points), so there is no horizon at which the naive "
         "benchmark would be the better choice.")

    table(document,
          ["Aggregation level", "Series", "WMAPE"],
          [
              ["product x store", "45,295", "43.6%"],
              ["product", "24,752", "35.7%"],
              ["store", "4,857", "18.7%"],
              ["category", "854", "15.0%"],
              ["region", "610", "11.8%"],
              ["**total**", "122", "**9.6%**"],
          ],
          widths=[2.4, 1.6, 1.6])
    para(document,
         "Error falls by a factor of 4.5 from SKU to total. Bottom-up aggregation is exactly "
         "coherent by construction, so nothing is reconciled - this table quantifies the "
         "price of that choice, and answers 'should I trust the regional number more than "
         "the SKU number?' with a magnitude.")

    callout(document, "Reported rather than smoothed.",
            "The horizon gradient is shallow - 43.6% at h1-3 against 44.7% at h57-90, only "
            "1.1 points over three months. At 1.25x the noise floor there are only ~9 points "
            "of learnable signal in total, so a steep curve is arithmetically unavailable. "
            "That makes the gradient weak evidence about join correctness on this dataset, "
            "and the leakage tests carry that argument instead.")

    callout(document, "The parameter trap worth remembering.",
            "XGBoost first scored 82.9% WMAPE at +58% bias and looked far worse than "
            "LightGBM. The cause was semantics: `min_child_weight` sums Hessians, and under "
            "`count:poisson` the Hessian is approximately mu, so the parameter scales with "
            "the target level. LightGBM's `min_child_samples` counts rows. Setting both to "
            "50 gave XGBoost ~38x less regularisation. Scaled to 50 x mean(y) it scores "
            "43.8%. Reporting the unscaled run would have put a confident and completely "
            "false line in the comparison table.", DANGER)

    heading(document, "Behaviour beyond the end of the data", 3)
    para(document,
         "The calendar, promotion schedule and price plan all end 2025-12-31, so a 90-day "
         "forecast is only fully informed from as-of 2025-10-02. Requests past that are "
         "**refused** with a recoverable error naming the latest workable as-of.")
    code(document, """
insufficient_data: a 90-day horizon from 2025-12-01 reaches 2026-03-01, but the
calendar, promotion schedule and price plan end 2025-12-31. Forecasting past that
would mean assuming no promotions are planned, which biases those days low. The
latest as-of that supports a 90-day horizon is 2025-10-02.
""")
    para(document,
         "The alternative - assume no promotion runs and carry the last price forward - "
         "produces a number that is systematically low and indistinguishable from a real "
         "forecast.", italic=True)
    page_break(document)


def part_crosscutting(document: Document) -> None:
    heading(document, "Part IV  -  Cross-cutting concerns", 1)

    heading(document, "Leakage: defence in depth", 2)
    para(document,
         "Leakage has the worst signal-to-noise of any failure mode in applied ML: it never "
         "raises, never warns, and makes every number look better. The only defence is to "
         "assert the absence of something, repeatedly, from several angles.")
    figure(document, "11_leakage_defence",
           "Figure 10  -  Five independent layers, each catching what the others miss.")

    callout(document, "The test that makes the rest meaningful.",
            "A test that has never failed proves nothing about its ability to detect "
            "anything. So the suite plants the exact bug the design prevents - "
            "`horizon_features_from_target=True` sources origin features from the target row "
            "- and asserts the detector fires. Without that, the mutation test is "
            "unfalsifiable.")

    table(document,
          ["Test", "Asserts"],
          [
              ["Mutation (T1)", "Corrupting all OBSERVED data after a cutoff leaves training "
                                "features byte-identical"],
              ["Falsifiability (T2)", "**T1 fails** when the bug is deliberately planted"],
              ["Arithmetic (T3)", "`lag_7` at origin `t` equals units at `t-7`, reconstructed "
                                  "by hand from the source panel"],
              ["Horizon monotonicity (T4)", "Error over long horizons exceeds short ones"],
              ["Noise floor (T5)", "Nothing scores below the irreducible 35% WMAPE"],
              ["Train/serve (T6)", "Both feature paths produce identical vectors"],
              ["Embargo (T7)", "No training target lands inside an evaluation fold"],
          ],
          widths=[1.9, 4.1])
    para(document,
         "T6 has already earned its place twice: it caught festival columns going missing "
         "over short serving windows, and categorical dtypes being inferred per frame.",
         italic=True)

    heading(document, "Testing philosophy", 2)
    bullets(document, [
        "**Reconstruct expectations by hand.** A test that calls the implementation to compute "
        "its own expected value cannot detect a bug in that implementation.",
        "**Prefer behavioural assertions.** 'Error grows with horizon' survives a refactor that "
        "renames every feature; a test naming columns does not.",
        "**Contextualise, do not hard-code.** Comparing WMAPE against a measured noise floor is "
        "a statement about the model. Comparing it against a threshold that happened to pass "
        "is a statement about how the fixture is tuned.",
        "**Test the failure path first.** Most of what matters about a service is how it behaves "
        "when something is missing, and that is the path least likely to be exercised by hand.",
    ])
    table(document,
          ["Marker", "Covers"],
          [
              ["`unit`", "Fast, isolated, no I/O"],
              ["`data`", "Synthetic generation and dataset invariants"],
              ["`features`", "Feature engineering and the feature repository"],
              ["`models`", "ML model behaviour"],
              ["`leakage`", "Point-in-time correctness - the property Steps 4-11 depend on"],
              ["`statistical`", "Relationship recovery against known ground truth"],
              ["`integration`", "API and on-disk artifacts"],
          ],
          widths=[1.4, 4.6])

    heading(document, "Reproducibility", 2)
    para(document,
         "Every training run records the seed, dataset version, feature version, code "
         "commit, split boundaries, hyper-parameters, and a **config fingerprint** - one "
         "hash over the entire configuration. Two runs sharing a fingerprint used the same "
         "setup; two that differ are not comparable, and the difference is discoverable "
         "rather than argued about.")
    callout(document, "A guard learned expensively.",
            "The evaluation report is written **before** MLflow tracking, and a tracking "
            "failure never fails the run. A three-hour Step 4 run was lost when MLflow "
            "rejected its own default store after every model had already been fitted. "
            "Bookkeeping must never destroy the thing it is keeping books on.", DANGER)

    heading(document, "Serving and the agent contract", 2)
    figure(document, "10_serving_path",
           "Figure 11  -  What the agent sees, and what it is deliberately never told.")
    para(document,
         "The agent should be able to state a number with its uncertainty and re-plan around "
         "a refusal. It should **not** need to know that LightGBM exists, where the parquet "
         "lives, how features are built, or what MLflow is - and none of that appears in the "
         "tool schema.")
    callout(document, "On confidence.",
            "`confidence` is measured interval coverage, or it is absent. There is no third "
            "option where a plausible-looking number appears because the field exists. A "
            "fabricated 0.89 costs nothing to emit and means nothing.")
    page_break(document)


def part_operations(document: Document) -> None:
    heading(document, "Part V  -  Operations", 1)

    heading(document, "Setup", 2)
    code(document, """
git clone <repository>
cd agentic-revenue-intelligence

.\\tasks.ps1 setup          # uv sync --all-extras, copies .env.example to .env
uv run ari config          # show resolved settings
uv run ari health          # probe every dependency
""")

    heading(document, "Generating data", 2)
    code(document, """
uv run ari generate-data --profile dev --seed 42 --validate

# Profiles:
#   smoke  - 40 products x 30 stores, used by CI and the test suite
#   dev    - 300 x 200, the working dataset
#   stress - larger, for performance work
""")

    heading(document, "Training the models", 2)
    code(document, """
# Step 4 - baseline sales
uv run python scripts/train_baseline.py --profile dev --seed 42
uv run python scripts/train_baseline.py --sample-pairs 200 --no-backtest --no-track

# Step 5 - demand forecasting
uv run python scripts/train_forecast.py --smoke --no-track    # <60s, correctness
uv run python scripts/train_forecast.py --seed 42             # ~8 min, the deliverable
uv run python scripts/train_forecast.py --full                # hours, declared
""")
    callout(document, "Sampled runs write elsewhere.",
            "A sampled run persists to `models/*_sampled/`, never to the directory the "
            "service loads from. Step 4 learned this the expensive way: a 400-pair "
            "verification run silently replaced a model trained on the full panel, and the "
            "only visible trace was a smaller calibration count in the metadata sidecar.")

    heading(document, "Quality gates", 2)
    code(document, """
.\\tasks.ps1 check       # ruff + mypy + pytest
.\\tasks.ps1 lint
.\\tasks.ps1 typecheck
.\\tasks.ps1 test
.\\tasks.ps1 security    # bandit
""")

    heading(document, "Configuration reference", 2)
    table(document,
          ["File", "Controls"],
          [
              ["`.env`", "Environment selection, credentials, paths. Never committed"],
              ["`configs/data/*.yaml`", "Dataset profiles - scale, causal parameters, confounders"],
              ["`configs/features/features.yaml`", "Which feature groups each model dataset uses"],
              ["`configs/models/forecasting.yaml`", "Horizons, sampling, split, embargo, buckets"],
              ["`pyproject.toml`", "Dependencies, ruff, mypy, pytest, bandit, coverage"],
          ],
          widths=[2.3, 3.7])

    para(document, "Key environment variables:")
    code(document, """
APP__ENVIRONMENT=local | databricks     # selects every seam implementation
DATA__PARQUET_ROOT=data/local/gold
ML__TRACKING_URI=sqlite:///data/local/mlflow.db
LLM__API_KEY=<secret>                   # never hard-coded, never logged
""")

    heading(document, "Databricks migration", 2)
    para(document,
         "Design only - nothing is implemented, per the project's staging rule. The full "
         "design lives in `docs/databricks_migration.md` and "
         "`docs/forecasting_databricks_migration.md`.")
    table(document,
          ["Moves unchanged", "Genuinely changes"],
          [
              ["Config, splits, conformal, evaluation, estimators, schemas, the whole tool "
               "contract",
               "Panel construction (pandas to PySpark), the horizon self-join (to a Delta "
               "join), the future scaffold, model loading (joblib to Unity Catalog), "
               "monitoring (to Lakehouse Monitoring)"],
          ],
          widths=[3.0, 3.0])
    callout(document, "What migration would not fix.",
            "The stockout-exclusion bias is a modelling decision, not a scale problem. The "
            "frozen competitor price is a data-availability fact. Distributing the "
            "statistical models faster would not make them appropriate at product-store "
            "grain. These travel unchanged and are named so the migration is not oversold.")
    page_break(document)


def part_retrospective(document: Document) -> None:
    heading(document, "Part VI  -  Retrospective", 1)

    heading(document, "The defect catalogue", 2)
    para(document,
         "Every defect below was found by a test, a measurement, or a diagnostic that "
         "existed to be sceptical - not by code review. They are recorded because the "
         "pattern is more instructive than any individual bug.")

    heading(document, "Defects that produced wrong numbers silently", 3)
    table(document,
          ["Defect", "How it surfaced"],
          [
              ["Inventory replenishment ordered against on-hand rather than inventory "
               "position, so stockouts were impossible",
               "A dataset invariant test on days-of-cover"],
              ["A query returned exactly 100,000 of 6.7M rows with no warning",
               "A deliberate check that row counts matched expectations"],
              ["`revenue` reached the feature matrix, recovering the target exactly",
               "A blanket correlation check against the target"],
              ["Supply features taught the baseline that low stock predicts low demand",
               "The stockout recovery diagnostic against latent demand"],
              ["`sample_pairs=N` loaded ~6.9x N series via an accidental cross product",
               "Measuring the loaded series count against the request"],
              ["Sampling was not reproducible - DuckDB returns GROUP BY rows unordered",
               "A determinism test comparing two identical calls"],
              ["Categorical dtypes inferred per frame gave inconsistent integer codes",
               "XGBoost refused an unseen category; LightGBM had accepted it silently"],
              ["XGBoost mis-regularised by 38x due to Hessian-scaled `min_child_weight`",
               "A bias of +58% that could not be explained by overfitting"],
          ],
          widths=[3.2, 2.8])

    heading(document, "Defects in the diagnostics themselves", 3)
    para(document,
         "Three cases where the measuring instrument was wrong rather than the model. These "
         "are the most dangerous kind, because they produce confident, incorrect verdicts.")
    table(document,
          ["Defect", "Consequence"],
          [
              ["The stockout criterion compared the baseline to **latent** demand, ignoring "
               "that stockouts are endogenous and latent demand during one runs 1.57x normal",
               "All six candidates were disqualified. The criterion, not the models, was wrong"],
              ["`explain.py` grouped the 364-day seasonal anchor with `lag_1` under one "
               "'demand history' family",
               "The combined share rose with horizon and looked exactly like the leakage "
               "alarm the function documents. Two features with opposite horizon profiles "
               "had been averaged into one diagnostic"],
              ["The scalability claim said per-series ETS was 'hours per backtest fold'",
               "Measurement showed 0.25s per fit - about 25 minutes for a full pass. "
               "Expensive, not infeasible. The claim was corrected and the argument moved to "
               "where the evidence supports it"],
          ],
          widths=[3.2, 2.8])

    heading(document, "Process defects", 3)
    table(document,
          ["Defect", "Rule now followed"],
          [
              ["A three-hour training run lost to an MLflow store rejection raised after "
               "training completed",
               "Write the report before tracking; tracking failures never fail the run"],
              ["A sampled verification run silently overwrote a full-panel model",
               "Sampled runs write to their own directory"],
              ["Regex patching of Python source deleted a method parameter",
               "Never regex source code - use the editor"],
              ["`New-Item -Force` truncated existing `__init__.py` files",
               "Check existence before creating"],
              ["`Set-Content` corrupted a UTF-8 notebook",
               "Use the file-writing tools or Python for content, never `Set-Content`"],
          ],
          widths=[3.2, 2.8])

    heading(document, "Known limitations", 2)
    table(document,
          ["Limitation", "Consequence"],
          [
              ["Excluding stockout targets biases the forecast low",
               "Stockouts are endogenous, so this removes part of the high-demand tail. "
               "Measured against latent demand and reported"],
              ["No inventory signal in either model",
               "Both forecast **demand**. A replenishment planner wants a different model"],
              ["Competitor prices frozen at the forecast origin",
               "A competitor move inside the horizon is invisible"],
              ["No cannibalisation or halo modelling",
               "A promotion on one SKU distorts its substitutes' numbers, unattributed"],
              ["Promotional baseline cannot be point-validated",
               "The store-level promo responsiveness multiplier is latent and unpublished, "
               "so validation there is directional only"],
              ["Cannot forecast past 2025-12-31",
               "A data limitation. Refused explicitly rather than faked"],
              ["Conformal assumes exchangeability",
               "A trend violates it mildly; coverage is measured per bucket so a shortfall "
               "surfaces"],
              ["Trained on synthetic data",
               "Real retail has structure this simulator does not reproduce"],
          ],
          widths=[2.6, 3.4])

    heading(document, "Roadmap", 2)
    figure(document, "12_roadmap",
           "Figure 12  -  Each model becomes a tool; the agent layer arrives last.")
    table(document,
          ["Steps", "Capability", "Depends on"],
          [
              ["6", "Promotion uplift - causal, not predictive", "Step 4 baseline"],
              ["7", "Price elasticity - using the cost instrument", "Steps 2-3"],
              ["8", "Cross-price elasticity and cannibalisation", "Step 7"],
              ["9-11", "Price optimisation, trade promo optimisation, scenario simulation",
               "Steps 5-8"],
              ["12-15", "Tool layer, prompts, memory, guardrails", "All models"],
              ["16-20", "Supervisor, critic, LangGraph workflow, evaluation", "Step 13+"],
          ],
          widths=[0.9, 3.1, 2.0])

    heading(document, "Talking points", 2)
    para(document,
         "The parts of this project worth discussing are rarely the model scores. They are "
         "the decisions where a plausible alternative was rejected for a stated reason.")
    bullets(document, [
        "**Why a structural causal model rather than sampled data.** Because it makes every "
        "later claim checkable. Almost no real project can validate a demand model against "
        "true demand; this one can, and that is what caught the supply-feature defect.",
        "**Why availability classes rather than a global as-of clamp.** Because the naive rule "
        "is wrong in a way that looks safe. Clamping the promotion calendar would have made "
        "forecasting impossible while appearing more rigorous.",
        "**Why the noise floor is reported alongside every accuracy figure.** Because a bare "
        "WMAPE is uninterpretable. 40% reads as poor until you know 35% is unlearnable; then "
        "it reads as 1.15x optimal.",
        "**Why the leakage suite plants a bug.** Because a test that has never failed is "
        "indistinguishable from a test that does nothing.",
        "**Why the forecaster refuses to forecast past the calendar.** Because the alternative "
        "produces a number that is systematically low and looks exactly like a real forecast.",
        "**Why `confidence` is measured coverage or absent.** Because a fabricated confidence "
        "score is the easiest thing in the system to emit and the hardest for a downstream "
        "consumer to detect.",
        "**Why XGBoost's first result was thrown out.** Because a 38x regularisation difference "
        "is a configuration error, and reporting it as a model comparison would have been a "
        "confident falsehood.",
    ])

    callout(document, "The through-line.",
            "Every mechanism in this platform exists because a specific plausible mistake "
            "would otherwise be invisible. That is the argument for the whole design: not "
            "that the models are accurate, but that when they are wrong, something says so.")


def build() -> None:
    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)

    for section in document.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    cover(document)

    heading(document, "Contents", 1)
    toc(document)
    page_break(document)

    how_to_read(document)
    part_orientation(document)
    part_foundations(document)
    part_models(document)
    part_crosscutting(document)
    part_operations(document)
    part_retrospective(document)

    document.save(OUTPUT)
    size_kb = OUTPUT.stat().st_size / 1024
    print(f"Written {OUTPUT}  ({size_kb:,.0f} KB)")


if __name__ == "__main__":
    build()
